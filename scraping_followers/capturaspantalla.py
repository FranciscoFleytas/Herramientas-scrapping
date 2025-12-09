import time
import os
import json
import random
import sys
import subprocess
import webbrowser
import shutil  # Nueva librería para eliminar carpetas
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Mm
from PIL import Image

# --- CONFIGURACIÓN ---
BATCH_SIZE = 50  # Cantidad de perfiles por tanda
MOBILE_UA = "Mozilla/5.0 (iPhone; CPU iPhone OS 16_6 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.6 Mobile/15E148 Safari/604.1"
WINDOW_WIDTH = 414
WINDOW_HEIGHT = 896
GEMINI_URL = "https://gemini.google.com/app/d483a5624d68b286?hl=es&hl=es"

# --- RUTAS ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CUENTAS_FILE = os.path.join(SCRIPT_DIR, 'cuentas.json')
INPUT_TXT_FILE = os.path.join(SCRIPT_DIR, 'perfiles.txt')
BASE_OUTPUT_DIR = os.path.join(SCRIPT_DIR, "SCREENSHOTS_MOBILE")

def setup_mobile_driver():
    options = uc.ChromeOptions()
    options.add_argument(f'--user-agent={MOBILE_UA}')
    options.add_argument("--disable-notifications")
    options.add_argument("--disable-popup-blocking")
    options.page_load_strategy = 'eager'
    
    driver = uc.Chrome(options=options, use_subprocess=True, version_main=142)
    driver.set_window_size(WINDOW_WIDTH, WINDOW_HEIGHT)
    return driver

def load_account():
    if os.path.exists(CUENTAS_FILE):
        with open(CUENTAS_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if isinstance(data, list) and len(data) > 0:
                return data[0] 
    return None

def load_target_profiles():
    if not os.path.exists(INPUT_TXT_FILE):
        return []
    with open(INPUT_TXT_FILE, 'r', encoding='utf-8') as f:
        return [line.strip() for line in f if line.strip()]

def login_with_cookie(driver, account):
    print(f"--- Logueando como {account['user']}... ---")
    driver.get("https://www.instagram.com/404")
    time.sleep(1) 
    driver.add_cookie({
        'name': 'sessionid',
        'value': account['sessionid'],
        'domain': '.instagram.com',
        'path': '/',
        'secure': True,
        'httpOnly': True
    })
    driver.get("https://www.instagram.com/")
    try:
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//main")))
    except:
        time.sleep(2)

def crear_carpeta_sesion():
    if not os.path.exists(BASE_OUTPUT_DIR):
        os.makedirs(BASE_OUTPUT_DIR)
    i = 1
    while True:
        folder_name = os.path.join(BASE_OUTPUT_DIR, f"capturas {i}")
        if not os.path.exists(folder_name):
            os.makedirs(folder_name)
            return folder_name
        i += 1

def capture_profile_views(driver, username, save_folder):
    print(f" > {username}", end=" ")
    driver.get(f"https://www.instagram.com/{username}/")
    
    captured_images = []
    user_folder = os.path.join(save_folder, username)
    if not os.path.exists(user_folder):
        os.makedirs(user_folder)

    try:
        WebDriverWait(driver, 8).until(EC.presence_of_element_located((By.TAG_NAME, "header")))
        
        path_header = os.path.join(user_folder, "01_header.png")
        driver.save_screenshot(path_header)
        captured_images.append(path_header)
        print(f"| F1", end=" ")

        driver.execute_script(f"window.scrollBy(0, {WINDOW_HEIGHT * 0.8});") 
        time.sleep(0.8)
        path_grid = os.path.join(user_folder, "02_grid.png")
        driver.save_screenshot(path_grid)
        captured_images.append(path_grid)
        print(f"| F2", end=" ")

        driver.execute_script(f"window.scrollBy(0, {WINDOW_HEIGHT * 0.8});")
        time.sleep(0.8)
        path_grid_2 = os.path.join(user_folder, "03_grid_more.png")
        driver.save_screenshot(path_grid_2)
        captured_images.append(path_grid_2)
        print(f"| F3 | OK")

    except Exception as e:
        print(f"| ERROR: {e}")

    return captured_images

def comprimir_imagen(ruta_png, calidad=65):
    try:
        ruta_jpg = ruta_png.replace(".png", ".jpg")
        if os.path.exists(ruta_jpg): return ruta_jpg
        with Image.open(ruta_png) as img:
            img = img.convert("RGB")
            img.save(ruta_jpg, "JPEG", quality=calidad, optimize=True)
        return ruta_jpg
    except:
        return ruta_png

def generate_word_report(session_folder, data_map, batch_number):
    print(f"\n--- Generando Word Parte {batch_number} ---")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(297)
    section.page_height = Mm(420)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    doc.add_heading(f'Reporte de Perfiles - Parte {batch_number}', 0)

    for username, images in data_map.items():
        doc.add_heading(f"Perfil: {username}", level=1)
        if len(images) > 0:
            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            for i, img_path in enumerate(images):
                if i < 3 and os.path.exists(img_path):
                    final_path = comprimir_imagen(img_path, calidad=65)
                    cell = table.cell(0, i)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(final_path, width=Mm(85))
        doc.add_page_break()
    
    report_path = os.path.join(session_folder, f"Reporte_Parte_{batch_number}.docx")
    doc.save(report_path)
    print(f"✅ DOC GUARDADO: {report_path}")

def generar_txt_tanda(session_folder, profiles_list, batch_number):
    """Genera el TXT correspondiente a la tanda actual."""
    txt_filename = f"textos_analizados_parte_{batch_number}.txt"
    txt_path = os.path.join(session_folder, txt_filename)
    try:
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f"LISTA DE PERFILES - PARTE {batch_number}\n")
            f.write("====================================\n")
            for perfil in profiles_list:
                f.write(f"{perfil}\n")
        print(f"✅ TXT GUARDADO: {txt_filename}")
    except Exception as e:
        print(f"Error TXT: {e}")

def limpiar_archivos_temporales(session_folder, profiles_list):
    """Elimina las carpetas de imágenes de los usuarios ya procesados para ahorrar espacio."""
    print("🧹 Limpiando imágenes temporales de esta tanda...", end=" ")
    count = 0
    for username in profiles_list:
        user_folder = os.path.join(session_folder, username)
        if os.path.exists(user_folder):
            try:
                shutil.rmtree(user_folder) # Elimina carpeta y contenido
                count += 1
            except Exception as e:
                print(f"[x] Error borrando {username}: {e}")
    print(f"({count} carpetas eliminadas)")

def main():
    account = load_account()
    if not account: return

    target_profiles = load_target_profiles()
    if not target_profiles: return

    current_session_folder = crear_carpeta_sesion()
    driver = setup_mobile_driver()
    
    try:
        login_with_cookie(driver, account)
        
        # Iterar por lotes (Chunks)
        for i in range(0, len(target_profiles), BATCH_SIZE):
            current_batch = target_profiles[i : i + BATCH_SIZE]
            batch_num = (i // BATCH_SIZE) + 1
            
            print(f"\n>>> INICIANDO LOTE {batch_num} (Perfiles {i+1} a {i+len(current_batch)})")
            
            batch_data = {} 
            for user in current_batch:
                images = capture_profile_views(driver, user, current_session_folder)
                batch_data[user] = images
                time.sleep(random.uniform(1.0, 2.0))
            
            # 1. Generar Word (las imágenes se insertan en el docx)
            generate_word_report(current_session_folder, batch_data, batch_num)
            
            # 2. Generar TXT de la tanda
            generar_txt_tanda(current_session_folder, batch_data.keys(), batch_num)
            
            # 3. Limpieza: Eliminar las carpetas de imágenes de ESTE lote
            # (El Word ya está guardado, así que es seguro borrarlas)
            limpiar_archivos_temporales(current_session_folder, batch_data.keys())
            
            del batch_data
        
        print("\n--- Abriendo carpeta y Gemini ---")
        if sys.platform.startswith('linux'):
            subprocess.Popen(['xdg-open', current_session_folder])
        webbrowser.open(GEMINI_URL)
            
    except Exception as e:
        print(f"Error global: {e}")
    finally:
        driver.quit()
        print("\n" + "="*40 + "\nPROCESO FINALIZADO\n" + "="*40)

if __name__ == "__main__":
    main()